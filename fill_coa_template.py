from docx import Document

def fill_coa_template(header_data, general_rows, micro_rows):
    doc = Document("templates/Internal_COA_Template.docx")

    # Fill header placeholders
    for p in doc.paragraphs:
        for key, value in header_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)

    # Main Table = first table
    main_table = doc.tables[0]
    for row_data in general_rows:
        row = main_table.add_row().cells
        row[0].text = row_data["Characteristic"]
        row[1].text = row_data["Standard"]
        row[2].text = row_data["Result"]
        row[3].text = row_data["Method"]

    # Micro Table = second table
    micro_table = doc.tables[1]
    for row_data in micro_rows:
        row = micro_table.add_row().cells
        row[0].text = row_data["Characteristic"]
        row[1].text = row_data["Standard"]
        row[2].text = row_data["Result"]
        row[3].text = row_data["Method"]

    doc.save("outputs/Generated_COA.docx")
