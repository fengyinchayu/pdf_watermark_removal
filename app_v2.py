import streamlit as st
import pdfplumber
from docx import Document
import os
import zipfile
from datetime import datetime
import re

# ==============================
# Utility Functions
# ==============================

def is_micro_test(name):
    keywords = [
        "Total Aerobic",
        "Mold",
        "Yeast",
        "Coliform",
        "Salmonella",
        "E.coli",
        "E. coli",
        "Staphylococcus"
    ]
    return any(k.lower() in name.lower() for k in keywords)


# -------------------------------------------------------
# Label aliases: maps known label variants from PDFs
# to our internal session_state keys.
# -------------------------------------------------------
HEADER_LABEL_MAP = {
    # Product name
    "product name":             "product_name",
    "product":                  "product_name",
    "item name":                "product_name",
    "item":                     "product_name",

    # Brand
    "brand":                    "brand",
    "manufacturer":             "brand",

    # Lot / Batch number
    "lot no":                   "lot_no",
    "lot no.":                  "lot_no",
    "lot number":               "lot_no",
    "lot#":                     "lot_no",
    "batch no":                 "lot_no",
    "batch no.":                "lot_no",
    "batch number":             "lot_no",
    "batch#":                   "lot_no",
    "batch":                    "lot_no",

    # Quantity / Batch size
    "quantity":                 "quantity",
    "qty":                      "quantity",
    "batch size":               "quantity",
    "size":                     "quantity",
    "net weight":               "quantity",
    "net wt":                   "quantity",
    "net wt.":                  "quantity",
    "total quantity":           "quantity",

    # Manufacturing date
    "mfg. date":                "mfg_date",
    "mfg.date":                 "mfg_date",
    "mfg date":                 "mfg_date",
    "mfg":                      "mfg_date",
    "manufacturing date":       "mfg_date",
    "manufacture date":         "mfg_date",
    "manufactured date":        "mfg_date",
    "date of manufacture":      "mfg_date",
    "date of manufacturing":    "mfg_date",
    "production date":          "mfg_date",
    "produced date":            "mfg_date",

    # Expiry date
    "expiry date":              "expiry_date",
    "expiry":                   "expiry_date",
    "expire date":              "expiry_date",
    "expiration date":          "expiry_date",
    "expiration":               "expiry_date",
    "exp. date":                "expiry_date",
    "exp.date":                 "expiry_date",
    "exp date":                 "expiry_date",
    "exp":                      "expiry_date",
    "best before":              "expiry_date",
    "use by":                   "expiry_date",

    # Shelf life
    "shelf life":               "shelf_life",
    "shelflife":                "shelf_life",

    # Plant part
    "plant part":               "plant_part",
    "part":                     "plant_part",
    "part used":                "plant_part",
    "parts used":               "plant_part",
    "used part":                "plant_part",

    # Latin / botanical name
    "plant latin name":         "latin_name",
    "latin name":               "latin_name",
    "botanical name":           "latin_name",
    "botanical":                "latin_name",
    "scientific name":          "latin_name",
    "plant name":               "latin_name",

    # Country of origin
    "country of origin":        "origin",
    "country":                  "origin",
    "origin":                   "origin",
    "place of origin":          "origin",
    "source country":           "origin",

    # Solvent
    "solvent":                  "solvent",
    "extraction solvent":       "solvent",
    "solvent of extraction":    "solvent",
}

FIELD_DEFAULTS = {
    "product_name": "",
    "brand":        "Skyherb®",
    "lot_no":       "",
    "quantity":     "",
    "mfg_date":     "",
    "shelf_life":   "3 years",
    "plant_part":   "",
    "latin_name":   "",
    "origin":       "China",
    "solvent":      "",
    "expiry_date":  "",
}


def normalise(text):
    """Lowercase + collapse whitespace for fuzzy label matching."""
    return re.sub(r"\s+", " ", text.strip().lower())


def extract_header_info(pdf_file):
    found = {}
    weight_re = re.compile(r"^\d[\d,. ]*\s*(kg|g|lbs?|lb)", re.IGNORECASE)

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    cells = [c.strip() if c else "" for c in row]
                    for i in range(0, len(cells) - 1, 2):
                        label = normalise(cells[i])
                        value = cells[i + 1].strip()
                        if not value:
                            continue
                        if label in HEADER_LABEL_MAP:
                            key = HEADER_LABEL_MAP[label]
                            if key not in found:
                                found[key] = value
                        elif "quantity" not in found and weight_re.match(value):
                            found["quantity"] = value

    return found


def extract_spec_tables(pdf_file):
    general_rows = []
    micro_rows = []

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if not row or not row[0]:
                        continue
                    if len(row) >= 3:
                        row_data = {
                            "Characteristic": row[0].strip(),
                            "Specification":  row[1].strip() if row[1] else "",
                            "Method":         row[2].strip() if row[2] else "",
                        }
                        if is_micro_test(row_data["Characteristic"]):
                            micro_rows.append(row_data)
                        else:
                            general_rows.append(row_data)

    return general_rows, micro_rows


def extract_coa_tables(pdf_file):
    general_rows = []
    micro_rows = []

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if not row or not row[0]:
                        continue
                    if len(row) >= 4:
                        row_data = {
                            "Characteristic": row[0].strip(),
                            "Standard":       row[1].strip() if row[1] else "",
                            "Result":         row[2].strip() if row[2] else "",
                            "Method":         row[3].strip() if row[3] else "",
                        }
                        if is_micro_test(row_data["Characteristic"]):
                            micro_rows.append(row_data)
                        else:
                            general_rows.append(row_data)

    return general_rows, micro_rows


def replace_placeholders_in_doc(doc, header_data):
    def replace_in_paragraph(paragraph, data):
        full_text = "".join(run.text for run in paragraph.runs)
        replaced = False
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
                replaced = True
        if replaced:
            for i, run in enumerate(paragraph.runs):
                run.text = full_text if i == 0 else ""

    for p in doc.paragraphs:
        replace_in_paragraph(p, header_data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, header_data)


def find_table_by_header(doc, keywords):
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(
            cell.text.strip().lower() for cell in table.rows[0].cells
        )
        if any(kw.lower() in first_row_text for kw in keywords):
            return table
    return None


def safe_write_row(cells, values):
    for i, val in enumerate(values):
        if i < len(cells):
            cells[i].text = val


def fill_spec_template(header_data, general_rows, micro_rows):
    doc = Document("templates/Internal_SPEC_Template.docx")
    replace_placeholders_in_doc(doc, header_data)

    main_table  = find_table_by_header(doc, ["characteristic", "specification", "method"])
    micro_table = find_table_by_header(doc, ["microbiological", "microbiology", "aerobic", "yeast", "mold", "coliform"])

    if main_table is None:
        if doc.tables:
            main_table = doc.tables[0]
            st.warning("Could not identify the SPEC main table by header — using the first table as fallback.")
        else:
            st.error("No tables found in the SPEC template.")
            return None

    for row_data in general_rows:
        safe_write_row(
            main_table.add_row().cells,
            [row_data.get("Characteristic", ""),
             row_data.get("Specification", ""),
             row_data.get("Method", "")]
        )

    if micro_table is not None:
        for row_data in micro_rows:
            safe_write_row(
                micro_table.add_row().cells,
                [row_data.get("Characteristic", ""),
                 row_data.get("Specification", ""),
                 row_data.get("Method", "")]
            )
    elif micro_rows:
        st.warning("No microbiological table found in the SPEC template — micro rows were skipped.")

    output_path = "outputs/Generated_SPEC.docx"
    doc.save(output_path)
    return output_path


def fill_coa_template(header_data, general_rows, micro_rows):
    doc = Document("templates/Internal_COA_Template.docx")
    replace_placeholders_in_doc(doc, header_data)

    main_table  = find_table_by_header(doc, ["characteristic", "standard", "result", "method"])
    micro_table = find_table_by_header(doc, ["microbiological", "microbiology", "aerobic", "yeast", "mold", "coliform"])

    if main_table is None:
        if doc.tables:
            main_table = doc.tables[0]
            st.warning("Could not identify the COA main table by header — using the first table as fallback.")
        else:
            st.error("No tables found in the COA template.")
            return None

    for row_data in general_rows:
        safe_write_row(
            main_table.add_row().cells,
            [row_data.get("Characteristic", ""),
             row_data.get("Standard", ""),
             row_data.get("Result", ""),
             row_data.get("Method", "")]
        )

    if micro_table is not None:
        for row_data in micro_rows:
            safe_write_row(
                micro_table.add_row().cells,
                [row_data.get("Characteristic", ""),
                 row_data.get("Standard", ""),
                 row_data.get("Result", ""),
                 row_data.get("Method", "")]
            )
    elif micro_rows:
        st.warning("No microbiological table found in the COA template — micro rows were skipped.")

    output_path = "outputs/Generated_COA.docx"
    doc.save(output_path)
    return output_path


# ==============================
# Watermark / Header / Scan Functions
# ==============================

from pypdf import PdfReader, PdfWriter
from pypdf.generic import ContentStream, NameObject, DecodedStreamObject
import io
from PIL import Image as PILImage

HEADER_IMAGE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SN_Header.png")


def _extract_image_thumbnail(xobj):
    try:
        width  = int(xobj.get("/Width",  0))
        height = int(xobj.get("/Height", 0))
        if width == 0 or height == 0:
            return None

        data = xobj.get_data()

        filter_val = xobj.get("/Filter", "")
        filters = [str(f) for f in filter_val] if hasattr(filter_val, "__iter__") and not isinstance(filter_val, str) else [str(filter_val)]

        if "/DCTDecode" in filters:
            img = PILImage.open(io.BytesIO(data))
        else:
            cs = xobj.get("/ColorSpace", "/DeviceRGB")
            cs_str = str(cs[0] if isinstance(cs, list) else cs)
            if "RGB"  in cs_str: mode, channels = "RGB",  3
            elif "CMYK" in cs_str: mode, channels = "CMYK", 4
            else:                  mode, channels = "L",    1
            expected = width * height * channels
            if len(data) < expected:
                return None
            img = PILImage.frombytes(mode, (width, height), data[:expected])

        img.thumbnail((180, 180), PILImage.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def scan_xobjects(pdf_bytes):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    seen   = {}

    for page_num, page in enumerate(reader.pages):
        content = page.get_contents()
        if content is None:
            continue

        try:
            content_obj = ContentStream(content, reader)
        except Exception:
            continue

        do_names = set()
        for operands, operator in content_obj.operations:
            if operator == b"Do" and operands:
                do_names.add(str(operands[0]))

        if not do_names:
            continue

        resources = page.get("/Resources", {})
        xobj_dict = resources.get("/XObject", {})

        for name_obj, xobj_ref in xobj_dict.items():
            name = str(name_obj)
            if name not in do_names:
                continue

            try:
                xobj = xobj_ref.get_object() if hasattr(xobj_ref, "get_object") else xobj_ref
            except Exception:
                continue

            subtype = str(xobj.get("/Subtype", "/Unknown")).lstrip("/")

            if name not in seen:
                thumbnail = _extract_image_thumbnail(xobj) if subtype == "Image" else None
                seen[name] = {
                    "name":      name,
                    "subtype":   subtype,
                    "pages":     [page_num + 1],
                    "thumbnail": thumbnail,
                }
            else:
                pg = page_num + 1
                if pg not in seen[name]["pages"]:
                    seen[name]["pages"].append(pg)

    return list(seen.values())


def remove_selected(input_pdf_bytes, output_path, names_to_remove, remove_gs=False):
    reader = PdfReader(io.BytesIO(input_pdf_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        content = page.get_contents()
        if content is None:
            writer.add_page(page)
            continue

        try:
            content_obj = ContentStream(content, reader)
        except Exception:
            writer.add_page(page)
            continue

        new_ops = []
        for operands, operator in content_obj.operations:
            if operator == b"Do" and operands:
                if str(operands[0]) in names_to_remove:
                    continue
            if operator == b"gs" and remove_gs:
                continue
            new_ops.append((operands, operator))

        content_obj.operations = new_ops
        new_stream = DecodedStreamObject()
        new_stream.set_data(content_obj.get_data())
        page[NameObject("/Contents")] = new_stream
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path


def make_header_overlay(page_width, page_height, header_img_path,
                        img_width_pt, top_margin_pt):
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c   = canvas.Canvas(buf, pagesize=(page_width, page_height))
    img = ImageReader(header_img_path)
    orig_w, orig_h = img.getSize()
    img_h_pt = img_width_pt * (orig_h / orig_w)
    c.drawImage(header_img_path, 0,
                page_height - top_margin_pt - img_h_pt,
                width=img_width_pt, height=img_h_pt, mask="auto")
    c.save()
    buf.seek(0)
    return buf.read()


def add_header_to_pdf(input_path, output_path, img_width_frac=0.98, top_margin_pt=10):
    reader = PdfReader(input_path)
    writer = PdfWriter()

    for page in reader.pages:
        pw = float(page.mediabox.width)
        ph = float(page.mediabox.height)
        overlay_bytes  = make_header_overlay(pw, ph, HEADER_IMAGE_PATH,
                                             pw * img_width_frac, top_margin_pt)
        overlay_page   = PdfReader(io.BytesIO(overlay_bytes)).pages[0]
        page.merge_page(overlay_page)
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path


# ==============================
# PDF Unlock Functions
# ==============================

def _try_unlock_reader(pdf_bytes: bytes, password: str) -> tuple[PdfReader | None, str]:
    """
    Attempt to open a PDF and decrypt it.
    Returns (reader, status_string).
    status: "not_encrypted" | "unlocked" | "wrong_password" | "error"
    """
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    except Exception as e:
        return None, f"error:{e}"

    if not reader.is_encrypted:
        return reader, "not_encrypted"

    # Try supplied password first, then empty string (owner lock / no user pw)
    passwords_to_try = [password] if password else []
    passwords_to_try.append("")  # always try blank as fallback

    for pw in passwords_to_try:
        try:
            result = reader.decrypt(pw)
            if result.value > 0:          # 1 = user pw, 2 = owner pw
                return reader, "unlocked"
        except Exception:
            pass

    return None, "wrong_password"


def unlock_pdf(pdf_bytes: bytes, password: str = "") -> tuple[bytes | None, str]:
    """
    Decrypt a PDF and return the raw bytes of the unlocked file.
    Returns (unlocked_bytes, status).
    status: "not_encrypted" | "unlocked" | "wrong_password" | "error:..."
    """
    reader, status = _try_unlock_reader(pdf_bytes, password)

    if status == "not_encrypted":
        # File is not encrypted — clone it as-is so the user still gets a clean copy
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue(), "not_encrypted"

    if status == "unlocked":
        writer = PdfWriter()
        writer.append_pages_from_reader(reader)
        # Explicitly clear encryption metadata
        writer._encrypt = None
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue(), "unlocked"

    return None, status


# ==============================
# Session State Initialisation
# ==============================

for field, default in FIELD_DEFAULTS.items():
    if field not in st.session_state:
        st.session_state[field] = default

os.makedirs("outputs", exist_ok=True)


# ==============================
# Sidebar Navigation
# ==============================

with st.sidebar:
    st.title("🧪 SN Tools")
    st.markdown("---")
    page = st.radio(
        "Navigate to",
        options=["📄 SPEC & COA Generator", "🚿 Watermark Remover", "🔓 PDF Unlocker"],
        label_visibility="collapsed"
    )
    st.markdown("---")
    st.caption("Skyherb® Internal Tools")


# ==============================
# PAGE 1 — SPEC & COA Generator
# ==============================

if page == "📄 SPEC & COA Generator":

    st.title("📄 SPEC & COA Generator")

    st.subheader("Upload Supplier Documents")

    col_up1, col_up2 = st.columns(2)
    with col_up1:
        spec_pdf = st.file_uploader("Upload Supplier SPEC (PDF)", type="pdf", key="spec_uploader")
    with col_up2:
        coa_pdf  = st.file_uploader("Upload Supplier COA (PDF)",  type="pdf", key="coa_uploader")

    extract_col1, extract_col2 = st.columns(2)

    with extract_col1:
        if spec_pdf and st.button("🔍 Extract Info from SPEC PDF"):
            with st.spinner("Extracting..."):
                extracted = extract_header_info(spec_pdf)
            if extracted:
                for key, value in extracted.items():
                    st.session_state[key] = value
                st.success(f"Extracted {len(extracted)} field(s) from SPEC PDF. Review and edit below.")
            else:
                st.warning("No recognisable header fields found in the SPEC PDF.")

    with extract_col2:
        if coa_pdf and st.button("🔍 Extract Info from COA PDF"):
            with st.spinner("Extracting..."):
                extracted = extract_header_info(coa_pdf)
            if extracted:
                for key, value in extracted.items():
                    st.session_state[key] = value
                st.success(f"Extracted {len(extracted)} field(s) from COA PDF. Review and edit below.")
            else:
                st.warning("No recognisable header fields found in the COA PDF.")

    st.divider()

    st.subheader("Product & Batch Information")
    st.caption("Fields are pre-filled from the PDF after clicking Extract. You can edit any value before generating.")

    col1, col2 = st.columns(2)

    with col1:
        product_name = st.text_input("Product Name",       value=st.session_state["product_name"], key="in_product_name")
        brand        = st.text_input("Brand",               value=st.session_state["brand"],        key="in_brand")
        lot_no       = st.text_input("Lot No.",             value=st.session_state["lot_no"],        key="in_lot_no")
        quantity     = st.text_input("Quantity",            value=st.session_state["quantity"],      key="in_quantity")
        mfg_date     = st.text_input("Mfg. Date",          value=st.session_state["mfg_date"],      key="in_mfg_date")
        expiry_date  = st.text_input("Expiry Date",        value=st.session_state["expiry_date"],   key="in_expiry_date")

    with col2:
        shelf_life   = st.text_input("Shelf Life",         value=st.session_state["shelf_life"],    key="in_shelf_life")
        plant_part   = st.text_input("Plant Part",         value=st.session_state["plant_part"],    key="in_plant_part")
        latin_name   = st.text_input("Plant Latin Name",   value=st.session_state["latin_name"],    key="in_latin_name")
        origin       = st.text_input("Country of Origin",  value=st.session_state["origin"],        key="in_origin")
        solvent      = st.text_input("Extraction Solvent", value=st.session_state["solvent"],       key="in_solvent")

    header_data = {
        "ProductName": product_name,
        "Brand":       brand,
        "Origin":      origin,
        "Solvent":     solvent,
        "PlantPart":   plant_part,
        "LatinName":   latin_name,
        "ShelfLife":   shelf_life,
        "LotNo":       lot_no,
        "Quantity":    quantity,
        "ManuDate":    mfg_date,
        "ExpiryDate":  expiry_date,
        "IssueDate":   datetime.today().strftime("%Y-%m-%d"),
    }

    st.divider()

    gen_col1, gen_col2 = st.columns(2)

    with gen_col1:
        if spec_pdf and st.button("⚙️ Generate Internal SPEC"):
            with st.spinner("Generating SPEC..."):
                general_rows, micro_rows = extract_spec_tables(spec_pdf)
                output_file = fill_spec_template(header_data, general_rows, micro_rows)
            if output_file:
                with open(output_file, "rb") as f:
                    st.download_button(
                        "⬇️ Download SPEC",
                        f,
                        file_name="Internal_SPEC.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    with gen_col2:
        if coa_pdf and st.button("⚙️ Generate Internal COA"):
            with st.spinner("Generating COA..."):
                general_rows, micro_rows = extract_coa_tables(coa_pdf)
                output_file = fill_coa_template(header_data, general_rows, micro_rows)
            if output_file:
                with open(output_file, "rb") as f:
                    st.download_button(
                        "⬇️ Download COA",
                        f,
                        file_name="Internal_COA.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


# ==============================
# PAGE 2 — Watermark Remover
# ==============================

elif page == "🚿 Watermark Remover":

    st.title("🚿 Watermark Remover")
    st.caption("Scan your PDFs to preview every embedded overlay, choose what to remove, then process.")

    if "wm_scan_results" not in st.session_state:
        st.session_state["wm_scan_results"] = []
    if "wm_remove_set" not in st.session_state:
        st.session_state["wm_remove_set"] = set()

    st.subheader("① Upload PDF(s)")

    uploaded_files = st.file_uploader(
        "Upload one or more PDF files",
        type="pdf",
        accept_multiple_files=True,
        key="wm_uploader"
    )

    if not uploaded_files:
        st.info("Upload one or more PDFs to get started.")
        st.stop()

    st.write(f"{len(uploaded_files)} file(s) uploaded.")

    st.divider()
    st.subheader("② Scan for Overlays & Images")

    if st.button("🔍 Scan PDFs for embedded elements"):
        combined = {}

        with st.spinner("Scanning..."):
            for uploaded in uploaded_files:
                try:
                    results = scan_xobjects(uploaded.read())
                    for item in results:
                        name = item["name"]
                        if name not in combined:
                            combined[name] = item.copy()
                        else:
                            combined[name]["pages"] = sorted(
                                set(combined[name]["pages"] + item["pages"])
                            )
                            if combined[name]["thumbnail"] is None and item["thumbnail"]:
                                combined[name]["thumbnail"] = item["thumbnail"]
                except Exception as e:
                    st.warning(f"Could not scan {uploaded.name}: {e}")

        st.session_state["wm_scan_results"] = list(combined.values())

        wm_hints = {"wm", "watermark", "mark", "stamp", "logo", "bg", "background"}
        auto_remove = set()
        for item in st.session_state["wm_scan_results"]:
            nm_lower = item["name"].lstrip("/").lower()
            if item["subtype"] == "Form":
                auto_remove.add(item["name"])
            elif any(h in nm_lower for h in wm_hints):
                auto_remove.add(item["name"])
        st.session_state["wm_remove_set"] = auto_remove

        if st.session_state["wm_scan_results"]:
            st.success(f"Found {len(st.session_state['wm_scan_results'])} embedded element(s). "
                       "Review them below and tick what to remove.")
        else:
            st.warning("No Do-operator XObjects found in these PDFs.")

    scan_results = st.session_state["wm_scan_results"]

    if scan_results:
        st.divider()
        st.subheader("③ Choose What to Remove")
        st.caption("Tick elements you want removed. Untick ones you want to keep. "
                   "Form XObjects and name-hinted items are pre-selected.")

        COLS = 4
        remove_set = st.session_state["wm_remove_set"].copy()

        for row_start in range(0, len(scan_results), COLS):
            row_items = scan_results[row_start : row_start + COLS]
            cols = st.columns(COLS)

            for col, item in zip(cols, row_items):
                with col:
                    name     = item["name"]
                    subtype  = item["subtype"]
                    pages    = item["pages"]
                    thumb    = item["thumbnail"]

                    if thumb:
                        st.image(thumb, use_container_width=True)
                    else:
                        if subtype == "Form":
                            icon, label = "📐", "Vector / Form"
                        else:
                            icon, label = "❓", subtype
                        st.markdown(
                            f"<div style='background:#f0f0f0;border-radius:6px;"
                            f"padding:28px 0;text-align:center;font-size:28px'>"
                            f"{icon}<br><small style='font-size:11px;color:#555'>{label}</small></div>",
                            unsafe_allow_html=True
                        )

                    page_str = (f"p.{pages[0]}" if len(pages) == 1
                                else f"p.{pages[0]}–{pages[-1]}" if pages == list(range(pages[0], pages[-1]+1))
                                else f"{len(pages)} pages")
                    st.caption(f"`{name}` · {subtype} · {page_str}")

                    checked = st.checkbox(
                        "Remove",
                        value=(name in remove_set),
                        key=f"chk_{name}"
                    )
                    if checked:
                        remove_set.add(name)
                    else:
                        remove_set.discard(name)

        st.session_state["wm_remove_set"] = remove_set

        st.divider()
        remove_gs = st.checkbox(
            "Also strip `gs` (graphic state) operators  ⚠️ may affect some page rendering",
            value=False,
            key="wm_remove_gs"
        )

        n_selected = len(st.session_state["wm_remove_set"])
        if n_selected:
            st.info(f"**{n_selected}** element(s) marked for removal: "
                    + ", ".join(f"`{n}`" for n in sorted(st.session_state["wm_remove_set"])))
        else:
            st.warning("Nothing selected — processing will leave PDFs unchanged.")

        st.divider()
        st.subheader("④ Company Header (optional)")

        header_exists = os.path.isfile(HEADER_IMAGE_PATH)

        if not header_exists:
            st.warning(f"⚠️ `SN_Header.png` not found next to `app.py`. "
                       "Place the file there to enable this option.")

        add_header = st.checkbox(
            "Stamp SN company header on every page",
            value=False,
            disabled=not header_exists,
            key="wm_add_header"
        )

        img_width_pct, top_margin = 98, 10
        if add_header and header_exists:
            st.image(HEADER_IMAGE_PATH,
                     caption="SN_Header.png — will be stamped on every page",
                     use_container_width=True)
            hcol1, hcol2 = st.columns(2)
            with hcol1:
                img_width_pct = st.slider(
                    "Header width (% of page width)",
                    min_value=20, max_value=100, value=98, step=1,
                    key="wm_header_width"
                )
            with hcol2:
                top_margin = st.number_input(
                    "Top margin (pt)",
                    min_value=0, max_value=100, value=10, step=1,
                    key="wm_top_margin"
                )

        st.divider()

        if st.button("⚙️ Process Files", type="primary"):
            names_to_remove = st.session_state["wm_remove_set"]
            results, errors = [], []
            progress = st.progress(0, text="Processing...")

            for i, uploaded in enumerate(uploaded_files):
                fname = uploaded.name
                try:
                    safe_name   = fname.replace(" ", "_")
                    output_name = f"SN-{safe_name}"
                    output_path = f"outputs/{output_name}"

                    remove_selected(
                        uploaded.read(),
                        output_path,
                        names_to_remove=names_to_remove,
                        remove_gs=remove_gs
                    )

                    if add_header and header_exists:
                        headed_path = output_path.replace(".pdf", "_headed.pdf")
                        add_header_to_pdf(output_path, headed_path,
                                          img_width_frac=img_width_pct / 100,
                                          top_margin_pt=top_margin)
                        os.replace(headed_path, output_path)

                    results.append((output_name, output_path))

                except Exception as e:
                    errors.append((fname, str(e)))

                progress.progress((i + 1) / len(uploaded_files),
                                  text=f"Processed {i + 1}/{len(uploaded_files)}")

            progress.empty()

            for name, err in errors:
                st.error(f"❌ {name}: {err}")

            if results:
                ops = []
                if names_to_remove: ops.append(f"{len(names_to_remove)} element(s) removed")
                if add_header and header_exists: ops.append("header added")
                st.success(f"✅ {len(results)} file(s) done" +
                           (f" ({', '.join(ops)})" if ops else "") + ".")
                st.divider()

                for output_name, output_path in results:
                    with open(output_path, "rb") as f:
                        st.download_button(
                            label=f"⬇️ {output_name}",
                            data=f,
                            file_name=output_name,
                            mime="application/pdf",
                            key=f"dl_{output_name}"
                        )


# ==============================
# PAGE 3 — PDF Unlocker
# ==============================

elif page == "🔓 PDF Unlocker":

    st.title("🔓 PDF Unlocker")
    st.caption(
        "Remove encryption and permission restrictions from PDFs. "
        "Supports owner-locked files (no password needed) and user-password-protected files."
    )

    # ── STEP 1 — Upload ───────────────────────────────────────────────────────
    st.subheader("① Upload PDF(s)")

    unlock_files = st.file_uploader(
        "Upload one or more locked PDF files",
        type="pdf",
        accept_multiple_files=True,
        key="unlock_uploader"
    )

    if not unlock_files:
        st.info("Upload one or more PDFs to get started.")
        st.stop()

    st.write(f"{len(unlock_files)} file(s) uploaded.")

    # ── STEP 2 — Password (optional) ─────────────────────────────────────────
    st.divider()
    st.subheader("② Password (optional)")
    st.caption(
        "Leave blank if the PDFs are owner-locked only (print/copy restrictions) "
        "with no user-open password. Enter a password only if the file asks for one to open."
    )

    unlock_password = st.text_input(
        "Password",
        type="password",
        placeholder="Leave blank if no open password is required",
        key="unlock_password"
    )

    same_pw_for_all = True
    if len(unlock_files) > 1:
        same_pw_for_all = st.checkbox(
            "Use the same password for all files",
            value=True,
            key="unlock_same_pw"
        )
        if not same_pw_for_all:
            st.info(
                "Per-file passwords are not supported in batch mode. "
                "Enter a shared password above, or process files one at a time."
            )

    # ── STEP 3 — Output naming ────────────────────────────────────────────────
    st.divider()
    st.subheader("③ Output Options")

    prefix = st.text_input(
        "Output filename prefix",
        value="unlocked-",
        help="This prefix is prepended to each original filename.",
        key="unlock_prefix"
    )

    # ── STEP 4 — Unlock ───────────────────────────────────────────────────────
    st.divider()

    if st.button("🔓 Unlock PDF(s)", type="primary"):

        results  = []   # list of (output_name, output_path, status)
        errors   = []   # list of (fname, reason)
        skipped  = []   # list of (fname, "not_encrypted")

        progress = st.progress(0, text="Unlocking...")

        for i, uploaded in enumerate(unlock_files):
            fname = uploaded.name
            try:
                pdf_bytes = uploaded.read()
                unlocked_bytes, status = unlock_pdf(pdf_bytes, password=unlock_password)

                safe_name   = fname.replace(" ", "_")
                output_name = f"{prefix}{safe_name}"
                output_path = f"outputs/{output_name}"

                if status in ("unlocked", "not_encrypted"):
                    with open(output_path, "wb") as f:
                        f.write(unlocked_bytes)
                    results.append((output_name, output_path, status))
                    if status == "not_encrypted":
                        skipped.append(fname)

                elif status == "wrong_password":
                    errors.append((fname, "Wrong or missing password — could not decrypt."))

                else:
                    # status starts with "error:"
                    errors.append((fname, status.replace("error:", "", 1)))

            except Exception as e:
                errors.append((fname, str(e)))

            progress.progress(
                (i + 1) / len(unlock_files),
                text=f"Processing {i + 1}/{len(unlock_files)}…"
            )

        progress.empty()

        # ── Results summary ────────────────────────────────────────────────
        unlocked_count  = len([r for r in results if r[2] == "unlocked"])
        passthru_count  = len([r for r in results if r[2] == "not_encrypted"])
        error_count     = len(errors)

        if unlocked_count:
            st.success(f"✅ {unlocked_count} file(s) successfully unlocked.")
        if passthru_count:
            st.info(
                f"ℹ️ {passthru_count} file(s) were not encrypted — "
                "copied through unchanged: "
                + ", ".join(f"`{n}`" for n in skipped)
            )
        for fname, reason in errors:
            st.error(f"❌ **{fname}**: {reason}")

        # ── Download buttons ───────────────────────────────────────────────
        if results:
            st.divider()
            st.subheader("⬇️ Download Unlocked Files")

            # Batch ZIP download (shown whenever there are 2+ files)
            if len(results) > 1:
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    for output_name, output_path, _ in results:
                        zf.write(output_path, arcname=output_name)
                zip_buf.seek(0)
                st.download_button(
                    label=f"⬇️ Download All ({len(results)} files) as ZIP",
                    data=zip_buf,
                    file_name="unlocked_pdfs.zip",
                    mime="application/zip",
                    key="dl_unlock_all_zip",
                    type="primary",
                    use_container_width=True,
                )
                st.caption("Or download files individually:")

            for output_name, output_path, status in results:
                label_suffix = " *(was not encrypted)*" if status == "not_encrypted" else ""
                with open(output_path, "rb") as f:
                    st.download_button(
                        label=f"⬇️ {output_name}{label_suffix}",
                        data=f,
                        file_name=output_name,
                        mime="application/pdf",
                        key=f"dl_unlock_{output_name}"
                    )

    # ── Helper note ───────────────────────────────────────────────────────────
    with st.expander("ℹ️ How does this work?"):
        st.markdown("""
**Two kinds of PDF locks:**

| Type | What it does | Password needed to unlock? |
|------|-------------|---------------------------|
| **User password** (open password) | Prevents opening the file | ✅ Yes — enter it above |
| **Owner password** (permissions lock) | Restricts printing, copying, editing | ❌ No — unlocked automatically |

**Batch processing:** Upload multiple files at once. All will be processed with the same password (if any).

**Output files** are prefixed with the text you set in ③ and are immediately available for download after processing.

> *This tool uses [pypdf](https://pypdf.readthedocs.io/) for decryption. It cannot brute-force unknown user passwords.*
        """)