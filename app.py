import streamlit as st
import pdfplumber
from docx import Document
import os
from datetime import datetime
import re

# ==============================
# Utility Functions
# ==============================

def is_micro_test(name):
    keywords = [
        "Total Aerobic",
        "Mold",
        "Yeast",
        "Coliform",
        "Salmonella",
        "E.coli",
        "E. coli",
        "Staphylococcus"
    ]
    return any(k.lower() in name.lower() for k in keywords)


# -------------------------------------------------------
# Label aliases: maps known label variants from PDFs
# to our internal session_state keys.
# -------------------------------------------------------
HEADER_LABEL_MAP = {
    # Product name
    "product name":             "product_name",
    "product":                  "product_name",
    "item name":                "product_name",
    "item":                     "product_name",

    # Brand
    "brand":                    "brand",
    "manufacturer":             "brand",

    # Lot / Batch number
    "lot no":                   "lot_no",
    "lot no.":                  "lot_no",
    "lot number":               "lot_no",
    "lot#":                     "lot_no",
    "batch no":                 "lot_no",
    "batch no.":                "lot_no",
    "batch number":             "lot_no",
    "batch#":                   "lot_no",
    "batch":                    "lot_no",

    # Quantity / Batch size
    "quantity":                 "quantity",
    "qty":                      "quantity",
    "batch size":               "quantity",
    "size":                     "quantity",
    "net weight":               "quantity",
    "net wt":                   "quantity",
    "net wt.":                  "quantity",
    "total quantity":           "quantity",

    # Manufacturing date
    "mfg. date":                "mfg_date",
    "mfg.date":                 "mfg_date",
    "mfg date":                 "mfg_date",
    "mfg":                      "mfg_date",
    "manufacturing date":       "mfg_date",
    "manufacture date":         "mfg_date",
    "manufactured date":        "mfg_date",
    "date of manufacture":      "mfg_date",
    "date of manufacturing":    "mfg_date",
    "production date":          "mfg_date",
    "produced date":            "mfg_date",

    # Expiry date
    "expiry date":              "expiry_date",
    "expiry":                   "expiry_date",
    "expire date":              "expiry_date",
    "expiration date":          "expiry_date",
    "expiration":               "expiry_date",
    "exp. date":                "expiry_date",
    "exp.date":                 "expiry_date",
    "exp date":                 "expiry_date",
    "exp":                      "expiry_date",
    "best before":              "expiry_date",
    "use by":                   "expiry_date",

    # Shelf life
    "shelf life":               "shelf_life",
    "shelflife":                "shelf_life",

    # Plant part
    "plant part":               "plant_part",
    "part":                     "plant_part",
    "part used":                "plant_part",
    "parts used":               "plant_part",
    "used part":                "plant_part",

    # Latin / botanical name
    "plant latin name":         "latin_name",
    "latin name":               "latin_name",
    "botanical name":           "latin_name",
    "botanical":                "latin_name",
    "scientific name":          "latin_name",
    "plant name":               "latin_name",

    # Country of origin
    "country of origin":        "origin",
    "country":                  "origin",
    "origin":                   "origin",
    "place of origin":          "origin",
    "source country":           "origin",

    # Solvent
    "solvent":                  "solvent",
    "extraction solvent":       "solvent",
    "solvent of extraction":    "solvent",
}

FIELD_DEFAULTS = {
    "product_name": "",
    "brand":        "Skyherb®",
    "lot_no":       "",
    "quantity":     "",
    "mfg_date":     "",
    "shelf_life":   "3 years",
    "plant_part":   "",
    "latin_name":   "",
    "origin":       "China",
    "solvent":      "",
    "expiry_date":  "",
}


def normalise(text):
    """Lowercase + collapse whitespace for fuzzy label matching."""
    return re.sub(r"\s+", " ", text.strip().lower())


def extract_header_info(pdf_file):
    """
    Scan every table in the PDF looking for label->value pairs.
    Handles both:
      2-column tables  (label | value)
      4-column tables  (label | value | label | value)

    Extra rule: if a cell value contains a weight pattern like
    25kg / 1000 kg / 500g and quantity has not been found yet,
    capture it regardless of the label text.

    Returns a dict keyed by our internal field names.
    """
    found = {}
    # Matches: 25kg / 25 kg / 1,000 kg / 2.5KG / 500g / 500 G
    weight_re = re.compile(r"^\d[\d,. ]*\s*(kg|g|lbs?|lb)", re.IGNORECASE)

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    # Clean cells
                    cells = [c.strip() if c else "" for c in row]

                    # Walk cells in label/value pairs: (0,1), (2,3), (4,5) ...
                    for i in range(0, len(cells) - 1, 2):
                        label = normalise(cells[i])
                        value = cells[i + 1].strip()

                        if not value:
                            continue

                        # Standard label-map lookup
                        if label in HEADER_LABEL_MAP:
                            key = HEADER_LABEL_MAP[label]
                            if key not in found:   # first occurrence wins
                                found[key] = value

                        # Fallback: value looks like a weight -> quantity
                        elif "quantity" not in found and weight_re.match(value):
                            found["quantity"] = value

    return found


def extract_spec_tables(pdf_file):
    general_rows = []
    micro_rows = []

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if not row or not row[0]:
                        continue
                    if len(row) >= 3:
                        row_data = {
                            "Characteristic": row[0].strip(),
                            "Specification":  row[1].strip() if row[1] else "",
                            "Method":         row[2].strip() if row[2] else "",
                        }
                        if is_micro_test(row_data["Characteristic"]):
                            micro_rows.append(row_data)
                        else:
                            general_rows.append(row_data)

    return general_rows, micro_rows


def extract_coa_tables(pdf_file):
    general_rows = []
    micro_rows = []

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if not row or not row[0]:
                        continue
                    if len(row) >= 4:
                        row_data = {
                            "Characteristic": row[0].strip(),
                            "Standard":       row[1].strip() if row[1] else "",
                            "Result":         row[2].strip() if row[2] else "",
                            "Method":         row[3].strip() if row[3] else "",
                        }
                        if is_micro_test(row_data["Characteristic"]):
                            micro_rows.append(row_data)
                        else:
                            general_rows.append(row_data)

    return general_rows, micro_rows


def replace_placeholders_in_doc(doc, header_data):
    """Replace {{key}} placeholders in all paragraphs and table cells."""

    def replace_in_paragraph(paragraph, data):
        full_text = "".join(run.text for run in paragraph.runs)
        replaced = False
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
                replaced = True
        if replaced:
            for i, run in enumerate(paragraph.runs):
                run.text = full_text if i == 0 else ""

    for p in doc.paragraphs:
        replace_in_paragraph(p, header_data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, header_data)


def find_table_by_header(doc, keywords):
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(
            cell.text.strip().lower() for cell in table.rows[0].cells
        )
        if any(kw.lower() in first_row_text for kw in keywords):
            return table
    return None


def safe_write_row(cells, values):
    for i, val in enumerate(values):
        if i < len(cells):
            cells[i].text = val


def fill_spec_template(header_data, general_rows, micro_rows):
    doc = Document("templates/Internal_SPEC_Template.docx")
    replace_placeholders_in_doc(doc, header_data)

    main_table  = find_table_by_header(doc, ["characteristic", "specification", "method"])
    micro_table = find_table_by_header(doc, ["microbiological", "microbiology", "aerobic", "yeast", "mold", "coliform"])

    if main_table is None:
        if doc.tables:
            main_table = doc.tables[0]
            st.warning("Could not identify the SPEC main table by header — using the first table as fallback.")
        else:
            st.error("No tables found in the SPEC template.")
            return None

    for row_data in general_rows:
        safe_write_row(
            main_table.add_row().cells,
            [row_data.get("Characteristic", ""),
             row_data.get("Specification", ""),
             row_data.get("Method", "")]
        )

    if micro_table is not None:
        for row_data in micro_rows:
            safe_write_row(
                micro_table.add_row().cells,
                [row_data.get("Characteristic", ""),
                 row_data.get("Specification", ""),
                 row_data.get("Method", "")]
            )
    elif micro_rows:
        st.warning("No microbiological table found in the SPEC template — micro rows were skipped.")

    output_path = "outputs/Generated_SPEC.docx"
    doc.save(output_path)
    return output_path


def fill_coa_template(header_data, general_rows, micro_rows):
    doc = Document("templates/Internal_COA_Template.docx")
    replace_placeholders_in_doc(doc, header_data)

    main_table  = find_table_by_header(doc, ["characteristic", "standard", "result", "method"])
    micro_table = find_table_by_header(doc, ["microbiological", "microbiology", "aerobic", "yeast", "mold", "coliform"])

    if main_table is None:
        if doc.tables:
            main_table = doc.tables[0]
            st.warning("Could not identify the COA main table by header — using the first table as fallback.")
        else:
            st.error("No tables found in the COA template.")
            return None

    for row_data in general_rows:
        safe_write_row(
            main_table.add_row().cells,
            [row_data.get("Characteristic", ""),
             row_data.get("Standard", ""),
             row_data.get("Result", ""),
             row_data.get("Method", "")]
        )

    if micro_table is not None:
        for row_data in micro_rows:
            safe_write_row(
                micro_table.add_row().cells,
                [row_data.get("Characteristic", ""),
                 row_data.get("Standard", ""),
                 row_data.get("Result", ""),
                 row_data.get("Method", "")]
            )
    elif micro_rows:
        st.warning("No microbiological table found in the COA template — micro rows were skipped.")

    output_path = "outputs/Generated_COA.docx"
    doc.save(output_path)
    return output_path


# ==============================
# Watermark Removal + Header Functions
# ==============================

from pypdf import PdfReader, PdfWriter
from pypdf.generic import ContentStream, NameObject, DecodedStreamObject
import io

HEADER_IMAGE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SN_Header.png")


def remove_watermark(input_pdf_bytes, output_path):
    reader = PdfReader(io.BytesIO(input_pdf_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        content = page.get_contents()
        content_obj = ContentStream(content, reader)

        new_operations = []
        for operands, operator in content_obj.operations:
            if operator in (b"Do", b"gs"):
                continue
            new_operations.append((operands, operator))

        content_obj.operations = new_operations

        new_stream = DecodedStreamObject()
        new_stream.set_data(content_obj.get_data())
        page[NameObject("/Contents")] = new_stream
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)

    return output_path


def make_header_overlay(page_width, page_height, header_img_path,
                        img_width_pt, top_margin_pt):
    """
    Build a single-page PDF (in memory) containing only the header image
    positioned at the top of a page of the given dimensions.
    Returns bytes of that PDF.
    """
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_width, page_height))

    img = ImageReader(header_img_path)
    orig_w, orig_h = img.getSize()
    aspect = orig_h / orig_w
    img_h_pt = img_width_pt * aspect

    x = 0
    y = page_height - top_margin_pt - img_h_pt

    c.drawImage(
        header_img_path,
        x, y,
        width=img_width_pt,
        height=img_h_pt,
        mask="auto"
    )
    c.save()
    buf.seek(0)
    return buf.read()


def add_header_to_pdf(input_path, output_path, img_width_frac=0.98, top_margin_pt=10):
    """
    Overlay the SN header image on every page of the PDF at input_path.
    img_width_frac: header image width as a fraction of the page width (0–1).
    top_margin_pt : gap between the top edge of the page and the header image, in points.
    Saves the result to output_path.
    """
    reader = PdfReader(input_path)
    writer = PdfWriter()

    for page in reader.pages:
        media_box = page.mediabox
        pw = float(media_box.width)
        ph = float(media_box.height)
        img_width_pt = pw * img_width_frac

        overlay_bytes = make_header_overlay(pw, ph, HEADER_IMAGE_PATH,
                                            img_width_pt, top_margin_pt)
        overlay_reader = PdfReader(io.BytesIO(overlay_bytes))
        overlay_page   = overlay_reader.pages[0]

        page.merge_page(overlay_page)
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)

    return output_path


# ==============================
# Session State Initialisation
# ==============================

for field, default in FIELD_DEFAULTS.items():
    if field not in st.session_state:
        st.session_state[field] = default

os.makedirs("outputs", exist_ok=True)


# ==============================
# Sidebar Navigation
# ==============================

with st.sidebar:
    st.title("🧪 SN Tools")
    st.markdown("---")
    page = st.radio(
        "Navigate to",
        options=["📄 SPEC & COA Generator", "🚿 Watermark Remover"],
        label_visibility="collapsed"
    )
    st.markdown("---")
    st.caption("Skyherb® Internal Tools")


# ==============================
# PAGE 1 — SPEC & COA Generator
# ==============================

if page == "📄 SPEC & COA Generator":

    st.title("📄 SPEC & COA Generator")

    # ── Upload section ────────────────────────────────────────────────────────
    st.subheader("Upload Supplier Documents")

    col_up1, col_up2 = st.columns(2)
    with col_up1:
        spec_pdf = st.file_uploader("Upload Supplier SPEC (PDF)", type="pdf", key="spec_uploader")
    with col_up2:
        coa_pdf  = st.file_uploader("Upload Supplier COA (PDF)",  type="pdf", key="coa_uploader")

    extract_col1, extract_col2 = st.columns(2)

    with extract_col1:
        if spec_pdf and st.button("🔍 Extract Info from SPEC PDF"):
            with st.spinner("Extracting..."):
                extracted = extract_header_info(spec_pdf)
            if extracted:
                for key, value in extracted.items():
                    st.session_state[key] = value
                st.success(f"Extracted {len(extracted)} field(s) from SPEC PDF. Review and edit below.")
            else:
                st.warning("No recognisable header fields found in the SPEC PDF.")

    with extract_col2:
        if coa_pdf and st.button("🔍 Extract Info from COA PDF"):
            with st.spinner("Extracting..."):
                extracted = extract_header_info(coa_pdf)
            if extracted:
                for key, value in extracted.items():
                    st.session_state[key] = value
                st.success(f"Extracted {len(extracted)} field(s) from COA PDF. Review and edit below.")
            else:
                st.warning("No recognisable header fields found in the COA PDF.")

    st.divider()

    # ── Editable product / batch fields ──────────────────────────────────────
    st.subheader("Product & Batch Information")
    st.caption("Fields are pre-filled from the PDF after clicking Extract. You can edit any value before generating.")

    col1, col2 = st.columns(2)

    with col1:
        product_name = st.text_input("Product Name",       value=st.session_state["product_name"], key="in_product_name")
        brand        = st.text_input("Brand",               value=st.session_state["brand"],        key="in_brand")
        lot_no       = st.text_input("Lot No.",             value=st.session_state["lot_no"],        key="in_lot_no")
        quantity     = st.text_input("Quantity",            value=st.session_state["quantity"],      key="in_quantity")
        mfg_date     = st.text_input("Mfg. Date",          value=st.session_state["mfg_date"],      key="in_mfg_date")
        expiry_date  = st.text_input("Expiry Date",        value=st.session_state["expiry_date"],   key="in_expiry_date")

    with col2:
        shelf_life   = st.text_input("Shelf Life",         value=st.session_state["shelf_life"],    key="in_shelf_life")
        plant_part   = st.text_input("Plant Part",         value=st.session_state["plant_part"],    key="in_plant_part")
        latin_name   = st.text_input("Plant Latin Name",   value=st.session_state["latin_name"],    key="in_latin_name")
        origin       = st.text_input("Country of Origin",  value=st.session_state["origin"],        key="in_origin")
        solvent      = st.text_input("Extraction Solvent", value=st.session_state["solvent"],       key="in_solvent")

    # ── Template placeholder mapping ──────────────────────────────────────────
    header_data = {
        "ProductName": product_name,
        "Brand":       brand,
        "Origin":      origin,
        "Solvent":     solvent,
        "PlantPart":   plant_part,
        "LatinName":   latin_name,
        "ShelfLife":   shelf_life,
        "LotNo":       lot_no,
        "Quantity":    quantity,
        "ManuDate":    mfg_date,
        "ExpiryDate":  expiry_date,
        "IssueDate":   datetime.today().strftime("%Y-%m-%d"),
    }

    st.divider()

    # ── Generate buttons ──────────────────────────────────────────────────────
    gen_col1, gen_col2 = st.columns(2)

    with gen_col1:
        if spec_pdf and st.button("⚙️ Generate Internal SPEC"):
            with st.spinner("Generating SPEC..."):
                general_rows, micro_rows = extract_spec_tables(spec_pdf)
                output_file = fill_spec_template(header_data, general_rows, micro_rows)
            if output_file:
                with open(output_file, "rb") as f:
                    st.download_button(
                        "⬇️ Download SPEC",
                        f,
                        file_name="Internal_SPEC.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    with gen_col2:
        if coa_pdf and st.button("⚙️ Generate Internal COA"):
            with st.spinner("Generating COA..."):
                general_rows, micro_rows = extract_coa_tables(coa_pdf)
                output_file = fill_coa_template(header_data, general_rows, micro_rows)
            if output_file:
                with open(output_file, "rb") as f:
                    st.download_button(
                        "⬇️ Download COA",
                        f,
                        file_name="Internal_COA.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


# ==============================
# PAGE 2 — Watermark Remover
# ==============================

elif page == "🚿 Watermark Remover":

    st.title("🚿 Watermark Remover")
    st.caption("Removes XObject overlays and graphic state operators typically used as PDF watermarks.")

    st.subheader("Upload PDF(s)")

    uploaded_files = st.file_uploader(
        "Upload one or more PDF files",
        type="pdf",
        accept_multiple_files=True,
        key="wm_uploader"
    )

    if uploaded_files:
        st.write(f"{len(uploaded_files)} file(s) ready.")

        st.divider()

        # ── Header overlay options ────────────────────────────────────────────
        st.subheader("Company Header")

        header_exists = os.path.isfile(HEADER_IMAGE_PATH)

        if not header_exists:
            st.warning(f"⚠️ Header image not found at `{HEADER_IMAGE_PATH}`. "
                       "Place `SN_Header.png` in the same folder as `app.py` to enable this option.")

        add_header = st.checkbox(
            "Add SN company header to each page",
            value=False,
            disabled=not header_exists,
            key="wm_add_header"
        )

        if add_header and header_exists:
            # Show a live preview of the header image
            st.image(HEADER_IMAGE_PATH, caption="SN_Header.png — will be stamped on every page", use_container_width=True)

            hcol1, hcol2 = st.columns(2)
            with hcol1:
                img_width_pct = st.slider(
                    "Header width (% of page width)",
                    min_value=20, max_value=100, value=98, step=1,
                    key="wm_header_width"
                )
            with hcol2:
                top_margin = st.number_input(
                    "Top margin (pt)",
                    min_value=0, max_value=100, value=10, step=1,
                    key="wm_top_margin"
                )

        st.divider()

        if st.button("⚙️ Process Files"):
            results = []
            errors  = []

            progress = st.progress(0, text="Processing...")

            for i, uploaded in enumerate(uploaded_files):
                fname = uploaded.name
                try:
                    safe_name   = fname.replace(" ", "_")
                    output_name = f"SN-{safe_name}"
                    output_path = f"outputs/{output_name}"

                    # Step 1 — remove watermark
                    remove_watermark(uploaded.read(), output_path)

                    # Step 2 — optionally stamp header on top
                    if add_header and header_exists:
                        headed_path = output_path.replace(".pdf", "_headed.pdf")
                        add_header_to_pdf(
                            output_path,
                            headed_path,
                            img_width_frac=img_width_pct / 100,
                            top_margin_pt=top_margin
                        )
                        os.replace(headed_path, output_path)  # overwrite, keep same filename

                    results.append((output_name, output_path))

                except Exception as e:
                    errors.append((fname, str(e)))

                progress.progress((i + 1) / len(uploaded_files),
                                  text=f"Processed {i + 1}/{len(uploaded_files)}")

            progress.empty()

            if errors:
                for name, err in errors:
                    st.error(f"❌ {name}: {err}")

            if results:
                label = "with watermark removed + header added" if add_header and header_exists                         else "with watermark removed"
                st.success(f"✅ {len(results)} file(s) processed ({label}).")
                st.divider()

                for output_name, output_path in results:
                    with open(output_path, "rb") as f:
                        st.download_button(
                            label=f"⬇️ {output_name}",
                            data=f,
                            file_name=output_name,
                            mime="application/pdf",
                            key=f"dl_{output_name}"
                        )