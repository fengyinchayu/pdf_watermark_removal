from docx import Document

def fill_spec_template(header_data, general_rows, micro_rows):
    doc = Document("templates/Internal_SPEC_Template.docx")

    for p in doc.paragraphs:
        for key, value in header_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)

    main_table = doc.tables[0]
    for row_data in general_rows:
        row = main_table.add_row().cells
        row[0].text = row_data["Characteristic"]
        row[1].text = row_data["Specification"]
        row[2].text = row_data["Method"]

    micro_table = doc.tables[1]
    for row_data in micro_rows:
        row = micro_table.add_row().cells
        row[0].text = row_data["Characteristic"]
        row[1].text = row_data["Specification"]
        row[2].text = row_data["Method"]

    doc.save("outputs/Generated_SPEC.docx")
